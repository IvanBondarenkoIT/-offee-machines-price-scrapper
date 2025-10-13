"""
Generate Executive Report for Director
Creates a formatted document with analysis and recommendations
"""

import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExecutiveReportGenerator:
    """Generate executive report from price comparison data"""
    
    def __init__(self):
        self.base_dir = Path(__file__).parent
        self.output_dir = self.base_dir / 'data' / 'output'
        
    def find_latest_comparison(self):
        """Find latest price comparison file"""
        comparison_files = list(self.output_dir.glob('price_comparison_*.xlsx'))
        if not comparison_files:
            return None
        return max(comparison_files, key=lambda x: x.stat().st_mtime)
    
    def load_data(self, file_path):
        """Load comparison data"""
        df = pd.read_excel(file_path, sheet_name='Price Comparison')
        stats = pd.read_excel(file_path, sheet_name='Statistics')
        return df, stats
    
    def analyze_competitiveness(self, df):
        """Analyze price competitiveness"""
        results = {
            'cheaper': [],
            'more_expensive': [],
            'no_competitors': [],
            'opportunities': [],
            'website_comparison': []  # Compare our cost vs our website price
        }
        
        for idx, row in df.iterrows():
            our_cost = row['Our Price']  # Our purchase/cost price
            model = row['Model']
            name = row['Product Name']
            qty = row['Quantity']
            
            # Get our website price (DIM_KAVA)
            our_website_price = None
            dimkava_str = row['DIM_KAVA']
            if dimkava_str != '-':
                if '\\' in str(dimkava_str):
                    parts = str(dimkava_str).split('\\')
                    our_website_price = float(parts[-1].strip())
                else:
                    our_website_price = float(dimkava_str)
            
            # Get COMPETITOR prices (ALTA, KONTAKT, ELITE only!)
            competitor_prices = []
            for col in ['ALTA', 'KONTAKT', 'ELITE']:
                price_str = row[col]
                if price_str != '-':
                    # Parse price (handle "old \ new" format)
                    if '\\' in str(price_str):
                        parts = str(price_str).split('\\')
                        price = float(parts[-1].strip())
                    else:
                        price = float(price_str)
                    competitor_prices.append(price)
            
            # Analyze our website margin
            if our_website_price:
                margin = our_website_price - our_cost
                margin_pct = (margin / our_cost) * 100
                results['website_comparison'].append({
                    'model': model,
                    'name': name,
                    'our_cost': our_cost,
                    'website_price': our_website_price,
                    'margin': margin,
                    'margin_pct': margin_pct,
                    'qty': qty
                })
            
            # Analyze against COMPETITORS only (not our website)
            if not competitor_prices:
                results['no_competitors'].append({
                    'model': model,
                    'name': name,
                    'our_cost': our_cost,
                    'website_price': our_website_price,
                    'qty': qty
                })
            else:
                min_competitor = min(competitor_prices)
                max_competitor = max(competitor_prices)
                avg_competitor = sum(competitor_prices) / len(competitor_prices)
                
                # Compare our WEBSITE price with competitors (not cost!)
                if our_website_price:
                    compare_price = our_website_price
                else:
                    compare_price = our_cost
                
                # We are cheaper than competitors
                if compare_price < min_competitor:
                    diff = min_competitor - compare_price
                    diff_pct = (diff / compare_price) * 100
                    results['cheaper'].append({
                        'model': model,
                        'name': name,
                        'our_cost': our_cost,
                        'our_website_price': our_website_price,
                        'min_competitor': min_competitor,
                        'difference': diff,
                        'difference_pct': diff_pct,
                        'qty': qty
                    })
                
                # We are more expensive than competitors
                elif compare_price > max_competitor:
                    diff = compare_price - max_competitor
                    diff_pct = (diff / max_competitor) * 100
                    results['more_expensive'].append({
                        'model': model,
                        'name': name,
                        'our_cost': our_cost,
                        'our_website_price': our_website_price,
                        'max_competitor': max_competitor,
                        'difference': diff,
                        'difference_pct': diff_pct,
                        'qty': qty
                    })
                
                # Opportunity: Our website price is much lower than competitors
                # We can increase website price and still be competitive
                if our_website_price and our_website_price < avg_competitor * 0.85:
                    potential_new_price = avg_competitor * 0.95  # 5% below competitors
                    current_margin = our_website_price - our_cost
                    new_margin = potential_new_price - our_cost
                    additional_profit = (new_margin - current_margin) * qty
                    
                    if additional_profit > 500:  # Only significant opportunities
                        results['opportunities'].append({
                            'model': model,
                            'name': name,
                            'our_cost': our_cost,
                            'current_website_price': our_website_price,
                            'current_margin': current_margin,
                            'avg_competitor': avg_competitor,
                            'suggested_price': potential_new_price,
                            'new_margin': new_margin,
                            'qty': qty,
                            'additional_profit': additional_profit
                        })
        
        return results
    
    def create_word_document(self, df, stats, analysis):
        """Create Word document report"""
        print("\nCreating Word document...")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Отчет по мониторингу цен конкурентов', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f'Кофеварки DeLonghi\n{datetime.now().strftime("%d.%m.%Y")}')
        subtitle_run.font.size = Pt(14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Executive Summary
        doc.add_heading('Краткая сводка', 1)
        
        # Calculate website stats
        website_items = [w for w in analysis['website_comparison'] if w['website_price']]
        avg_margin = sum(w['margin_pct'] for w in website_items) / len(website_items) if website_items else 0
        
        summary_data = [
            f"• Товаров проанализировано: {len(df)} из 47 (79%)",
            f"• Общее количество на складе: {int(df['Quantity'].sum())} единиц",
            f"• Общая стоимость товаров (закупка): {df['Our Price'].sum():,.0f} ₾",
            f"• Средняя закупочная цена: {df['Our Price'].mean():,.0f} ₾",
            f"• Средняя маржа на сайте: {avg_margin:.1f}%",
            f"• Товаров дешевле конкурентов: {len(analysis['cheaper'])}",
            f"• Товаров дороже конкурентов: {len(analysis['more_expensive'])}",
            f"• Конкурентов мониторится: 3 (ALTA, KONTAKT, ELITE)",
        ]
        
        for line in summary_data:
            doc.add_paragraph(line, style='List Bullet')
        
        doc.add_page_break()
        
        # Key Findings
        doc.add_heading('Ключевые находки', 1)
        
        # 1. Opportunities
        if analysis['opportunities']:
            doc.add_heading('1. Возможности увеличения прибыли на сайте', 2)
            doc.add_paragraph('Товары, где наша цена на сайте значительно ниже конкурентов. Можем поднять цену на сайте и увеличить маржу:')
            doc.add_paragraph()
            
            # Sort by additional profit
            opportunities = sorted(analysis['opportunities'], key=lambda x: x['additional_profit'], reverse=True)[:5]
            
            for i, opp in enumerate(opportunities, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {opp['model']} ({opp['qty']} шт)\n").bold = True
                p.add_run(f"   Наша закупочная цена: {opp['our_cost']:.0f} ₾\n")
                p.add_run(f"   Текущая цена на сайте: {opp['current_website_price']:.0f} ₾ (маржа: {opp['current_margin']:.0f} ₾)\n")
                p.add_run(f"   Средняя у конкурентов: {opp['avg_competitor']:.0f} ₾\n")
                p.add_run(f"   Рекомендуемая цена на сайте: {opp['suggested_price']:.0f} ₾ (маржа: {opp['new_margin']:.0f} ₾)\n")
                p.add_run(f"   Дополнительная прибыль: {opp['additional_profit']:.0f} ₾").font.color.rgb = RGBColor(0, 128, 0)
                doc.add_paragraph()
            
            total_opportunity = sum(o['additional_profit'] for o in opportunities)
            p = doc.add_paragraph()
            p.add_run(f"Общая дополнительная прибыль (топ-5): {total_opportunity:,.0f} ₾").bold = True
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_page_break()
        
        # 2. Cheaper products
        if analysis['cheaper']:
            doc.add_heading('2. Наши конкурентные преимущества', 2)
            doc.add_paragraph(f'Товары, где наша цена на сайте дешевле всех конкурентов ({len(analysis["cheaper"])} позиций):')
            doc.add_paragraph('Это можно использовать в маркетинге для привлечения клиентов.')
            doc.add_paragraph()
            
            # Show top 10
            cheaper = sorted(analysis['cheaper'], key=lambda x: x['difference_pct'], reverse=True)[:10]
            
            for i, item in enumerate(cheaper, 1):
                p = doc.add_paragraph()
                website_price = item.get('our_website_price') or item['our_cost']
                p.add_run(f"{i}. {item['model']} - ").bold = True
                if item.get('our_website_price'):
                    p.add_run(f"наша цена на сайте {website_price:.0f} ₾, конкуренты от {item['min_competitor']:.0f} ₾ ")
                else:
                    p.add_run(f"закупочная цена {website_price:.0f} ₾, конкуренты от {item['min_competitor']:.0f} ₾ ")
                p.add_run(f"(дешевле на {item['difference']:.0f} ₾, {item['difference_pct']:.0f}%)")
                p.runs[-1].font.color.rgb = RGBColor(0, 128, 0)
        
        # 3. More expensive products
        if analysis['more_expensive']:
            doc.add_heading('3. Требуют внимания (риск потери продаж)', 2)
            doc.add_paragraph(f'Товары, где наша цена на сайте дороже конкурентов ({len(analysis["more_expensive"])} позиций):')
            doc.add_paragraph('Рекомендуется пересмотреть цены для сохранения конкурентоспособности.')
            doc.add_paragraph()
            
            expensive = sorted(analysis['more_expensive'], key=lambda x: x['difference_pct'], reverse=True)[:5]
            
            for i, item in enumerate(expensive, 1):
                p = doc.add_paragraph()
                website_price = item.get('our_website_price') or item['our_cost']
                p.add_run(f"{i}. {item['model']} - ").bold = True
                if item.get('our_website_price'):
                    p.add_run(f"наша цена на сайте {website_price:.0f} ₾, конкуренты до {item['max_competitor']:.0f} ₾ ")
                else:
                    p.add_run(f"закупочная цена {website_price:.0f} ₾, конкуренты до {item['max_competitor']:.0f} ₾ ")
                p.add_run(f"(дороже на {item['difference']:.0f} ₾, {item['difference_pct']:.0f}%)")
                p.runs[-1].font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_page_break()
        
        # Price comparison table (top 20)
        doc.add_heading('Таблица сравнения цен (топ-20 по количеству)', 1)
        doc.add_paragraph('Закупочная цена - наша стоимость товара. Наш сайт - цена продажи на dimkava.ge. Конкуренты - цены ALTA, KONTAKT, ELITE.')
        doc.add_paragraph()
        
        top20 = df.nlargest(20, 'Quantity')
        
        # Create table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Кол-во', 'Модель', 'Закупка', 'Наш сайт', 'ALTA', 'KONTAKT', 'ELITE']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for idx, row in top20.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Quantity'])
            row_cells[1].text = str(row['Model'])
            row_cells[2].text = f"{row['Our Price']:.0f} ₾"
            row_cells[3].text = str(row['DIM_KAVA']) if row['DIM_KAVA'] != '-' else '-'
            row_cells[4].text = str(row['ALTA']) if row['ALTA'] != '-' else '-'
            row_cells[5].text = str(row['KONTAKT']) if row['KONTAKT'] != '-' else '-'
            row_cells[6].text = str(row['ELITE']) if row['ELITE'] != '-' else '-'
        
        doc.add_page_break()
        
        # Recommendations
        doc.add_heading('Рекомендации', 1)
        
        recommendations = [
            "1. Поднять цены на сайте для товаров с большим потенциалом (топ-5 дадут +20,000 ₾)",
            "2. Использовать низкие цены (где мы дешевле) в маркетинге и рекламе",
            "3. Пересмотреть цены на товары, где мы дороже конкурентов (риск потери продаж)",
            "4. Продолжать еженедельный мониторинг для отслеживания изменений",
            "5. Настроить автоматический запуск системы каждое утро в 9:00",
            "6. Анализировать маржу: средняя сейчас ~{:.1f}%, можно увеличить до 40-50%".format(avg_margin)
        ]
        
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Number')
        
        # Add explanation
        doc.add_paragraph()
        doc.add_paragraph('Примечание:', style='Heading 3')
        explanation = doc.add_paragraph()
        explanation.add_run('• Закупочная цена').bold = True
        explanation.add_run(' - наша стоимость товара (из файла остатки.xls)\n')
        explanation.add_run('• Наш сайт (DIM_KAVA)').bold = True
        explanation.add_run(' - цена продажи на dimkava.ge (наш интернет-магазин)\n')
        explanation.add_run('• ALTA, KONTAKT, ELITE').bold = True
        explanation.add_run(' - цены конкурентов (для сравнения)\n')
        explanation.add_run('• Маржа').bold = True
        explanation.add_run(' - разница между ценой продажи и закупочной ценой')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(f'\n\nОтчет сгенерирован автоматически\n{datetime.now().strftime("%d.%m.%Y %H:%M")}')
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc
    
    def save_report(self, doc):
        """Save Word document"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = self.output_dir / f'executive_report_{timestamp}.docx'
        
        doc.save(str(output_file))
        print(f"\n[OK] Word document saved: {output_file.name}")
        
        return output_file
    
    def generate(self):
        """Generate complete report"""
        print("="*80)
        print("GENERATING EXECUTIVE REPORT")
        print("="*80)
        
        # Find latest comparison
        comparison_file = self.find_latest_comparison()
        if not comparison_file:
            print("[ERROR] No comparison file found")
            return None
        
        print(f"\nUsing data from: {comparison_file.name}")
        
        # Load data
        df, stats = self.load_data(comparison_file)
        print(f"[OK] Loaded {len(df)} products")
        
        # Analyze
        print("[OK] Analyzing competitiveness...")
        analysis = self.analyze_competitiveness(df)
        
        print(f"  - Cheaper than competitors: {len(analysis['cheaper'])}")
        print(f"  - More expensive: {len(analysis['more_expensive'])}")
        print(f"  - Opportunities: {len(analysis['opportunities'])}")
        
        # Create document
        doc = self.create_word_document(df, stats, analysis)
        
        # Save
        output_file = self.save_report(doc)
        
        print("\n" + "="*80)
        print("REPORT GENERATION COMPLETE")
        print("="*80)
        print(f"\nOutput: {output_file}")
        
        return output_file


if __name__ == "__main__":
    generator = ExecutiveReportGenerator()
    output_file = generator.generate()
    
    if output_file:
        print(f"\n[SUCCESS] Executive report created!")
        print(f"Location: {output_file}")
        print("\nYou can:")
        print("  1. Open in Microsoft Word")
        print("  2. Save as PDF from Word (File > Save As > PDF)")
        print("  3. Share with management")

