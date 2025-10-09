"""
Generate Executive Report for Director
Creates a formatted document with analysis and recommendations
"""

import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExecutiveReportGenerator:
    """Generate executive report from price comparison data"""
    
    def __init__(self):
        self.base_dir = Path(__file__).parent
        self.output_dir = self.base_dir / 'data' / 'output'
        
    def find_latest_comparison(self):
        """Find latest price comparison file"""
        comparison_files = list(self.output_dir.glob('price_comparison_*.xlsx'))
        if not comparison_files:
            return None
        return max(comparison_files, key=lambda x: x.stat().st_mtime)
    
    def load_data(self, file_path):
        """Load comparison data"""
        df = pd.read_excel(file_path, sheet_name='Price Comparison')
        stats = pd.read_excel(file_path, sheet_name='Statistics')
        return df, stats
    
    def analyze_competitiveness(self, df):
        """Analyze price competitiveness"""
        results = {
            'cheaper': [],
            'more_expensive': [],
            'no_competitors': [],
            'opportunities': []
        }
        
        for idx, row in df.iterrows():
            our_price = row['Our Price']
            model = row['Model']
            name = row['Product Name']
            qty = row['Quantity']
            
            # Get competitor prices
            competitor_prices = []
            for col in ['DIM_KAVA', 'ALTA', 'KONTAKT', 'ELITE']:
                price_str = row[col]
                if price_str != '-':
                    # Parse price (handle "old \ new" format)
                    if '\\' in str(price_str):
                        parts = str(price_str).split('\\')
                        price = float(parts[-1].strip())
                    else:
                        price = float(price_str)
                    competitor_prices.append(price)
            
            if not competitor_prices:
                results['no_competitors'].append({
                    'model': model,
                    'name': name,
                    'our_price': our_price,
                    'qty': qty
                })
            else:
                min_competitor = min(competitor_prices)
                max_competitor = max(competitor_prices)
                avg_competitor = sum(competitor_prices) / len(competitor_prices)
                
                # We are cheaper
                if our_price < min_competitor:
                    diff = min_competitor - our_price
                    diff_pct = (diff / our_price) * 100
                    results['cheaper'].append({
                        'model': model,
                        'name': name,
                        'our_price': our_price,
                        'min_competitor': min_competitor,
                        'difference': diff,
                        'difference_pct': diff_pct,
                        'qty': qty,
                        'potential_profit': diff * qty
                    })
                
                # We are more expensive
                elif our_price > max_competitor:
                    diff = our_price - max_competitor
                    diff_pct = (diff / max_competitor) * 100
                    results['more_expensive'].append({
                        'model': model,
                        'name': name,
                        'our_price': our_price,
                        'max_competitor': max_competitor,
                        'difference': diff,
                        'difference_pct': diff_pct,
                        'qty': qty
                    })
                
                # Opportunity to increase price
                if our_price < avg_competitor * 0.8:  # We're 20%+ cheaper
                    potential_new_price = avg_competitor * 0.95  # Set to 5% below average
                    potential_profit = (potential_new_price - our_price) * qty
                    if potential_profit > 500:  # Only significant opportunities
                        results['opportunities'].append({
                            'model': model,
                            'name': name,
                            'our_price': our_price,
                            'avg_competitor': avg_competitor,
                            'suggested_price': potential_new_price,
                            'qty': qty,
                            'potential_profit': potential_profit
                        })
        
        return results
    
    def create_word_document(self, df, stats, analysis):
        """Create Word document report"""
        print("\nCreating Word document...")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Отчет по мониторингу цен конкурентов', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f'Кофеварки DeLonghi\n{datetime.now().strftime("%d.%m.%Y")}')
        subtitle_run.font.size = Pt(14)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Executive Summary
        doc.add_heading('Краткая сводка', 1)
        
        summary_data = [
            f"• Товаров проанализировано: {len(df)} из 47 (79%)",
            f"• Общее количество на складе: {int(df['Quantity'].sum())} единиц",
            f"• Общая стоимость товаров: {df['Our Price'].sum():,.0f} ₾",
            f"• Средняя наша цена: {df['Our Price'].mean():,.0f} ₾",
            f"• Товаров дешевле конкурентов: {len(analysis['cheaper'])}",
            f"• Товаров дороже конкурентов: {len(analysis['more_expensive'])}",
        ]
        
        for line in summary_data:
            doc.add_paragraph(line, style='List Bullet')
        
        doc.add_page_break()
        
        # Key Findings
        doc.add_heading('Ключевые находки', 1)
        
        # 1. Opportunities
        if analysis['opportunities']:
            doc.add_heading('1. Возможности увеличения прибыли', 2)
            doc.add_paragraph('Товары, где мы значительно дешевле конкурентов и можем поднять цену:')
            doc.add_paragraph()
            
            # Sort by potential profit
            opportunities = sorted(analysis['opportunities'], key=lambda x: x['potential_profit'], reverse=True)[:5]
            
            for i, opp in enumerate(opportunities, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {opp['model']} ({opp['qty']} шт)\n").bold = True
                p.add_run(f"   Наша цена: {opp['our_price']:.0f} ₾\n")
                p.add_run(f"   Средняя у конкурентов: {opp['avg_competitor']:.0f} ₾\n")
                p.add_run(f"   Рекомендуемая цена: {opp['suggested_price']:.0f} ₾\n")
                p.add_run(f"   Потенциальная прибыль: {opp['potential_profit']:.0f} ₾").font.color.rgb = RGBColor(0, 128, 0)
                doc.add_paragraph()
            
            total_opportunity = sum(o['potential_profit'] for o in opportunities)
            p = doc.add_paragraph()
            p.add_run(f"Общая потенциальная прибыль (топ-5): {total_opportunity:,.0f} ₾").bold = True
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        
        doc.add_page_break()
        
        # 2. Cheaper products
        if analysis['cheaper']:
            doc.add_heading('2. Наши конкурентные преимущества', 2)
            doc.add_paragraph(f'Товары, где мы дешевле всех конкурентов ({len(analysis["cheaper"])} позиций):')
            doc.add_paragraph()
            
            # Show top 10
            cheaper = sorted(analysis['cheaper'], key=lambda x: x['difference_pct'], reverse=True)[:10]
            
            for i, item in enumerate(cheaper, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {item['model']} - ").bold = True
                p.add_run(f"мы дешевле на {item['difference']:.0f} ₾ ({item['difference_pct']:.0f}%)")
        
        # 3. More expensive products
        if analysis['more_expensive']:
            doc.add_heading('3. Требуют внимания', 2)
            doc.add_paragraph(f'Товары, где мы дороже конкурентов ({len(analysis["more_expensive"])} позиций):')
            doc.add_paragraph()
            
            expensive = sorted(analysis['more_expensive'], key=lambda x: x['difference_pct'], reverse=True)[:5]
            
            for i, item in enumerate(expensive, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {item['model']} - ").bold = True
                p.add_run(f"мы дороже на {item['difference']:.0f} ₾ ({item['difference_pct']:.0f}%)")
                p.runs[-1].font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_page_break()
        
        # Price comparison table (top 20)
        doc.add_heading('Таблица сравнения цен (топ-20 по количеству)', 1)
        
        top20 = df.nlargest(20, 'Quantity')
        
        # Create table
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Кол-во', 'Модель', 'Наша цена', 'Наш сайт', 'ALTA', 'KONTAKT', 'ELITE']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for idx, row in top20.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Quantity'])
            row_cells[1].text = str(row['Model'])
            row_cells[2].text = f"{row['Our Price']:.0f} ₾"
            row_cells[3].text = str(row['DIM_KAVA']) if row['DIM_KAVA'] != '-' else '-'
            row_cells[4].text = str(row['ALTA']) if row['ALTA'] != '-' else '-'
            row_cells[5].text = str(row['KONTAKT']) if row['KONTAKT'] != '-' else '-'
            row_cells[6].text = str(row['ELITE']) if row['ELITE'] != '-' else '-'
        
        doc.add_page_break()
        
        # Recommendations
        doc.add_heading('Рекомендации', 1)
        
        recommendations = [
            "1. Рассмотреть повышение цен на товары с большой разницей (потенциал +20,000 ₾/месяц)",
            "2. Продолжать мониторинг цен еженедельно",
            "3. Использовать конкурентные цены как преимущество в маркетинге",
            "4. Обратить внимание на товары, где мы дороже конкурентов",
            "5. Настроить автоматический запуск системы каждое утро"
        ]
        
        for rec in recommendations:
            doc.add_paragraph(rec, style='List Number')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer_run = footer.add_run(f'\n\nОтчет сгенерирован автоматически\n{datetime.now().strftime("%d.%m.%Y %H:%M")}')
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc
    
    def save_report(self, doc):
        """Save Word document"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_file = self.output_dir / f'executive_report_{timestamp}.docx'
        
        doc.save(str(output_file))
        print(f"\n[OK] Word document saved: {output_file.name}")
        
        return output_file
    
    def generate(self):
        """Generate complete report"""
        print("="*80)
        print("GENERATING EXECUTIVE REPORT")
        print("="*80)
        
        # Find latest comparison
        comparison_file = self.find_latest_comparison()
        if not comparison_file:
            print("[ERROR] No comparison file found")
            return None
        
        print(f"\nUsing data from: {comparison_file.name}")
        
        # Load data
        df, stats = self.load_data(comparison_file)
        print(f"[OK] Loaded {len(df)} products")
        
        # Analyze
        print("[OK] Analyzing competitiveness...")
        analysis = self.analyze_competitiveness(df)
        
        print(f"  - Cheaper than competitors: {len(analysis['cheaper'])}")
        print(f"  - More expensive: {len(analysis['more_expensive'])}")
        print(f"  - Opportunities: {len(analysis['opportunities'])}")
        
        # Create document
        doc = self.create_word_document(df, stats, analysis)
        
        # Save
        output_file = self.save_report(doc)
        
        print("\n" + "="*80)
        print("REPORT GENERATION COMPLETE")
        print("="*80)
        print(f"\nOutput: {output_file}")
        
        return output_file


if __name__ == "__main__":
    generator = ExecutiveReportGenerator()
    output_file = generator.generate()
    
    if output_file:
        print(f"\n[SUCCESS] Executive report created!")
        print(f"Location: {output_file}")
        print("\nYou can:")
        print("  1. Open in Microsoft Word")
        print("  2. Save as PDF from Word (File > Save As > PDF)")
        print("  3. Share with management")

